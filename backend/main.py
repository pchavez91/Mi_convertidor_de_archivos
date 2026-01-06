from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os
import subprocess
import shutil
from pathlib import Path
import uuid
from typing import Optional
import aiofiles
import socket
import logging
import traceback
import asyncio

app = FastAPI(title="Convertidor de Archivos API")

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configurar CORS
# En producción, permitir el dominio del frontend
# En desarrollo, permitir localhost y red local
cors_origins = [
    "http://localhost:5173",
    "http://localhost:3000",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:3000",
]

# Agregar dominio del frontend si está configurado (Fly.io, Vercel, Netlify, etc.)
frontend_url = os.getenv("FRONTEND_URL", "")
if frontend_url:
    # Si hay múltiples URLs separadas por coma
    urls = [url.strip() for url in frontend_url.split(",")]
    for url in urls:
        cors_origins.append(url)
        # También agregar versión HTTPS si es HTTP
        if url.startswith("http://"):
            cors_origins.append(url.replace("http://", "https://"))
        # También agregar versión HTTP si es HTTPS (para desarrollo)
        elif url.startswith("https://"):
            cors_origins.append(url.replace("https://", "http://"))

# Agregar dominios conocidos de producción (hardcoded para seguridad adicional)
production_domains = [
    "https://todoconvertir.com",
    "https://www.todoconvertir.com",
    "http://todoconvertir.com",
    "http://www.todoconvertir.com",
]
cors_origins.extend(production_domains)

# Agregar IPs de red local comunes (solo en desarrollo)
try:
    hostname = socket.gethostname()
    local_ip = socket.gethostbyname(hostname)
    # Agregar IP local
    cors_origins.extend([
        f"http://{local_ip}:5173",
        f"http://{local_ip}:3000",
    ])
    # Agregar rangos comunes de red local
    ip_parts = local_ip.split('.')
    if len(ip_parts) == 4:
        base_ip = '.'.join(ip_parts[:3])
        # Permitir cualquier IP en el mismo segmento de red
        cors_origins.append(f"http://{base_ip}.*:5173")
        cors_origins.append(f"http://{base_ip}.*:3000")
except:
    pass

# Permitir todos los orígenes (ajustar en producción según necesidad)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En producción, cambiar a lista específica de dominios
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directorios para archivos temporales
# Detectar si estamos en un contenedor (Fly.io, Docker) o desarrollo local
if Path("/app").exists():
    # En contenedor (Fly.io, Docker)
    UPLOAD_DIR = Path("/app/backend/uploads")
    OUTPUT_DIR = Path("/app/backend/outputs")
elif Path("backend").exists() and not Path("/app").exists():
    # Desarrollo local desde la raíz del proyecto
    UPLOAD_DIR = Path("backend/uploads")
    OUTPUT_DIR = Path("backend/outputs")
else:
    # Desarrollo local desde backend/
    UPLOAD_DIR = Path("uploads")
    OUTPUT_DIR = Path("outputs")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Formatos soportados
AUDIO_FORMATS = ["mp3", "wav", "aac", "ogg", "flac", "m4a", "wma"]
VIDEO_FORMATS = ["mp4", "avi", "mov", "mkv", "webm", "flv", "wmv", "m4v"]
IMAGE_FORMATS = ["jpg", "jpeg", "png", "webp", "gif", "bmp", "ico", "tiff"]
DOCUMENT_FORMATS = ["pdf", "docx", "txt", "html", "md", "rtf", "odt"]

def get_file_type(filename: str) -> str:
    """Determina el tipo de archivo basado en la extensión"""
    ext = filename.split('.')[-1].lower()
    if ext in AUDIO_FORMATS:
        return "audio"
    elif ext in VIDEO_FORMATS:
        return "video"
    elif ext in IMAGE_FORMATS:
        return "image"
    elif ext in DOCUMENT_FORMATS:
        return "document"
    return "unknown"

def check_ffmpeg():
    """Verifica si FFmpeg está instalado"""
    try:
        subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False

def get_local_ip():
    """Obtiene la IP local del servidor"""
    try:
        # Conectarse a un servidor externo para obtener la IP local
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except:
        try:
            return socket.gethostbyname(socket.gethostname())
        except:
            return "127.0.0.1"

@app.get("/")
async def root():
    local_ip = get_local_ip()
    port = int(os.getenv("PORT", 8000))
    return {
        "message": "Convertidor de Archivos API",
        "status": "running",
        "local_ip": local_ip,
        "access_url": f"http://{local_ip}:{port}"
    }

@app.get("/formats")
async def get_formats():
    """Retorna los formatos soportados"""
    return {
        "audio": AUDIO_FORMATS,
        "video": VIDEO_FORMATS,
        "image": IMAGE_FORMATS,
        "document": DOCUMENT_FORMATS
    }

@app.post("/convert")
async def convert_file(
    file: UploadFile = File(...),
    output_format: str = Form(...)
):
    """Convierte un archivo al formato especificado"""
    
    if not output_format:
        raise HTTPException(status_code=400, detail="Formato de salida no especificado")
    
    output_format = output_format.lower()
    file_type = get_file_type(file.filename)
    
    if file_type == "unknown":
        raise HTTPException(status_code=400, detail="Formato de archivo no soportado")
    
    # Validar tamaño máximo del archivo (50 MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB en bytes
    file_size = 0
    
    # Leer el contenido del archivo y verificar tamaño
    content = await file.read()
    file_size = len(content)
    
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"El archivo es demasiado grande. Tamaño máximo permitido: 50 MB. Tu archivo: {file_size / (1024 * 1024):.2f} MB"
        )
    
    if file_size == 0:
        raise HTTPException(status_code=400, detail="El archivo está vacío")
    
    # Validar que el formato de salida sea diferente al de entrada
    input_ext = file.filename.split('.')[-1].lower()
    # Normalizar extensiones (jpg = jpeg)
    if input_ext == "jpeg":
        input_ext = "jpg"
    
    if input_ext == output_format:
        raise HTTPException(
            status_code=400, 
            detail=f"El archivo ya está en formato {output_format.upper()}. Por favor elige un formato de salida diferente."
        )
    
    # Generar nombres únicos
    file_id = str(uuid.uuid4())
    input_path = UPLOAD_DIR / f"{file_id}_{file.filename}"
    output_filename = f"{file_id}.{output_format}"
    output_path = OUTPUT_DIR / output_filename
    
    try:
        # Guardar archivo subido (ya tenemos el contenido en memoria)
        async with aiofiles.open(input_path, 'wb') as f:
            await f.write(content)
        
        # Realizar conversión según el tipo
        if file_type in ["audio", "video"]:
            if not check_ffmpeg():
                raise HTTPException(
                    status_code=500, 
                    detail="FFmpeg no está instalado. Por favor instálalo para convertir audio/video."
                )
            await convert_media(input_path, output_path, output_format)
        elif file_type == "image":
            await convert_image(input_path, output_path, output_format)
        elif file_type == "document":
            await convert_document(input_path, output_path, output_format)
        else:
            raise HTTPException(status_code=400, detail="Tipo de archivo no soportado")
        
        # Verificar que el archivo de salida existe
        if not output_path.exists():
            raise HTTPException(status_code=500, detail="Error en la conversión")
        
        return {
            "success": True,
            "download_url": f"/download/{output_filename}",
            "filename": output_filename
        }
    
    except HTTPException:
        # Re-lanzar HTTPException sin modificar
        if input_path.exists():
            input_path.unlink()
        if output_path.exists():
            output_path.unlink()
        raise
    except Exception as e:
        # Limpiar archivos en caso de error
        error_trace = traceback.format_exc()
        logger.error(f"Error en conversión: {str(e)}\n{error_trace}")
        if input_path.exists():
            input_path.unlink()
        if output_path.exists():
            output_path.unlink()
        raise HTTPException(status_code=500, detail=f"Error en la conversión: {str(e)}")

async def convert_media(input_path: Path, output_path: Path, output_format: str):
    """Convierte archivos de audio/video usando FFmpeg con optimizaciones"""
    
    # Construir comando base con optimizaciones
    cmd = [
        "ffmpeg",
        "-i", str(input_path),
        "-y",  # Sobrescribir archivo de salida
    ]
    
    # Optimizaciones generales para velocidad (priorizar velocidad sobre calidad máxima)
    cmd.extend([
        "-threads", "0",  # Usar todos los cores disponibles
    ])
    
    # Ajustes específicos para audio (optimizados para velocidad)
    if output_format in AUDIO_FORMATS:
        if output_format == "mp3":
            cmd.extend([
                "-codec:a", "libmp3lame",
                "-b:a", "192k",
                "-q:a", "4",  # Calidad buena pero más rápida (0-9, más alto = más rápido)
            ])
        elif output_format == "wav":
            cmd.extend([
                "-codec:a", "pcm_s16le",
            ])
        elif output_format == "aac":
            cmd.extend([
                "-codec:a", "aac",
                "-b:a", "192k",
                "-profile:a", "aac_low",  # Perfil más rápido
            ])
        elif output_format == "ogg":
            cmd.extend([
                "-codec:a", "libvorbis",
                "-q:a", "4",  # Calidad buena pero más rápida
            ])
        elif output_format == "flac":
            cmd.extend([
                "-codec:a", "flac",
                "-compression_level", "3",  # Nivel más bajo = más rápido
            ])
    else:
        # Para video, optimizaciones agresivas para velocidad
        if output_format == "webm":
            # WEBM con optimizaciones de velocidad
            cmd.extend([
                "-c:v", "libvpx-vp9",
                "-crf", "32",  # Calidad ligeramente menor pero más rápido
                "-b:v", "0",
                "-row-mt", "1",
                "-deadline", "realtime",  # Priorizar velocidad
                "-cpu-used", "4",  # Más rápido (0-8, más alto = más rápido)
                "-c:a", "libopus",
                "-b:a", "96k",  # Bitrate de audio más bajo para velocidad
            ])
        elif output_format == "mp4":
            cmd.extend([
                "-c:v", "libx264",
                "-preset", "veryfast",  # Preset más rápido que "fast"
                "-crf", "24",  # Calidad ligeramente menor pero más rápido
                "-tune", "fastdecode",  # Optimizar para decodificación rápida
                "-movflags", "+faststart",
                "-c:a", "aac",
                "-b:a", "96k",  # Bitrate de audio más bajo
            ])
        else:
            # Para otros formatos de video, usar configuración genérica optimizada
            cmd.extend([
                "-c:v", "libx264",
                "-preset", "veryfast",  # Preset más rápido
                "-crf", "24",
                "-tune", "fastdecode",
                "-c:a", "aac",
                "-b:a", "96k",
            ])
    
    cmd.append(str(output_path))
    
    # Ejecutar FFmpeg de forma asíncrona con timeout
    try:
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        
        # Esperar con timeout (30 minutos máximo)
        try:
            stdout, stderr = await asyncio.wait_for(
                process.communicate(),
                timeout=1800.0  # 30 minutos
            )
        except asyncio.TimeoutError:
            process.kill()
            await process.wait()
            raise Exception("La conversión excedió el tiempo máximo permitido (30 minutos)")
        
        if process.returncode != 0:
            error_msg = stderr.decode('utf-8', errors='ignore') if stderr else "Error desconocido"
            # Filtrar mensajes informativos comunes de FFmpeg
            error_lines = error_msg.split('\n')
            # Buscar líneas de error reales (no información de versión o configuración)
            real_errors = []
            skip_patterns = ['ffmpeg version', 'Copyright', 'built with', 'configuration:', 'libav', '--enable-']
            
            for line in error_lines:
                line_lower = line.lower()
                # Si la línea contiene palabras clave de error real
                if any(keyword in line_lower for keyword in ['error', 'failed', 'invalid', 'cannot', 'unable']):
                    real_errors.append(line)
                # Si no es información de configuración/versión, puede ser un error
                elif line.strip() and not any(skip in line_lower for skip in skip_patterns):
                    # Verificar si parece un mensaje de error
                    if ':' in line and ('error' in line_lower or 'failed' in line_lower):
                        real_errors.append(line)
            
            # Si encontramos errores reales, usarlos; sino usar el último mensaje significativo
            if real_errors:
                error_text = '\n'.join(real_errors[-5:])  # Últimos 5 errores
            else:
                # Si no hay errores claros, buscar las últimas líneas que no sean información de versión
                last_lines = [line for line in error_lines[-10:] if line.strip() and not any(skip in line.lower() for skip in skip_patterns)]
                error_text = '\n'.join(last_lines) if last_lines else error_msg[-500:]
            
            logger.error(f"FFmpeg error (returncode {process.returncode}): {error_text}")
            
            # Verificar si el archivo se creó a pesar del error
            if output_path.exists() and output_path.stat().st_size > 0:
                # Si el archivo existe y tiene tamaño, probablemente fue exitoso
                logger.info(f"Archivo de salida creado a pesar de returncode {process.returncode}, continuando...")
            else:
                # Error real, el archivo no se creó
                raise Exception(f"Error en FFmpeg: {error_text[:300]}")
        
    except FileNotFoundError:
        raise Exception("FFmpeg no está instalado o no está en el PATH")
    except Exception as e:
        logger.error(f"Error en convert_media: {str(e)}\n{traceback.format_exc()}")
        raise

async def convert_image(input_path: Path, output_path: Path, output_format: str):
    """Convierte imágenes usando Pillow"""
    from PIL import Image
    
    
    try:
        img = Image.open(input_path)
        
        # Manejar GIFs animados - tomar solo el primer frame
        if img.format == "GIF":
            try:
                # Verificar si es animado
                if hasattr(img, 'is_animated') and img.is_animated:
                    # Ir al primer frame y copiarlo
                    img.seek(0)
                    # Crear una copia del primer frame
                    first_frame = img.copy()
                    img = first_frame
            except Exception as e:
                logger.warning(f"Error manejando GIF animado, usando frame actual: {str(e)}")
                # Si hay error, simplemente usar la imagen tal como está
                pass
        
        # Convertir modos de color problemáticos a RGB/RGBA según necesidad
        # Esto debe hacerse antes de las conversiones específicas de formato
        
        # Convertir modos de color problemáticos
        if img.mode in ["P", "LA", "PA"]:
            # P = Palette, LA = Luminance + Alpha, PA = Palette + Alpha
            if img.mode == "P":
                # Verificar si tiene transparencia
                if "transparency" in img.info:
                    img = img.convert("RGBA")
                else:
                    img = img.convert("RGB")
            elif img.mode in ["LA", "PA"]:
                img = img.convert("RGBA")
        
        # Convertir RGBA a RGB si es necesario (para formatos que no soportan transparencia)
        if output_format in ["jpg", "jpeg"]:
            # JPG no soporta transparencia, siempre convertir a RGB
            if img.mode == "RGBA":
                # Crear fondo blanco y pegar la imagen con transparencia
                rgb_img = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "RGBA":
                    # Usar el canal alpha como máscara
                    alpha = img.split()[3] if len(img.split()) > 3 else None
                    rgb_img.paste(img, mask=alpha)
                else:
                    rgb_img.paste(img)
                img = rgb_img
            elif img.mode not in ["RGB", "L"]:
                # Convertir otros modos a RGB
                img = img.convert("RGB")
        elif output_format in ["bmp", "webp"]:
            # BMP y WEBP pueden manejar RGBA, pero mejor convertir a RGB si no hay transparencia
            if img.mode == "RGBA" and output_format == "bmp":
                # BMP no soporta transparencia bien, convertir a RGB
                rgb_img = Image.new("RGB", img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[3] if len(img.split()) > 3 else None)
                img = rgb_img
        elif output_format == "png":
            # PNG soporta transparencia, mantener RGBA si existe
            if img.mode not in ["RGB", "RGBA", "L", "LA", "P"]:
                img = img.convert("RGBA")
        elif output_format == "gif":
            # GIF puede tener transparencia
            if img.mode not in ["RGB", "RGBA", "P", "L"]:
                img = img.convert("RGB")
        
        # Determinar el formato para Pillow
        format_map = {
            "jpg": "JPEG",
            "jpeg": "JPEG",
            "png": "PNG",
            "gif": "GIF",
            "bmp": "BMP",
            "webp": "WEBP",
            "tiff": "TIFF",
            "ico": "ICO"
        }
        
        pillow_format = format_map.get(output_format.lower(), output_format.upper())
        
        # Opciones de guardado según el formato
        save_kwargs = {}
        if pillow_format == "JPEG":
            save_kwargs["quality"] = 95
            save_kwargs["optimize"] = True
        elif pillow_format == "WEBP":
            save_kwargs["quality"] = 90
        elif pillow_format == "PNG":
            save_kwargs["optimize"] = True
        
        # Guardar en el nuevo formato
        img.save(output_path, format=pillow_format, **save_kwargs)
        
    except Exception as e:
        logger.error(f"Error convirtiendo imagen: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Error al convertir imagen: {str(e)}"
        )

def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extrae texto de un archivo PDF"""
    try:
        import PyPDF2
        text_content = []
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="El archivo PDF está vacío o no tiene páginas"
                )
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
        
        if not text_content:
            raise HTTPException(
                status_code=400,
                detail="El archivo PDF no contiene texto extraíble. Puede ser un PDF escaneado (imagen) o estar protegido."
            )
        
        return "\n\n".join(text_content)
    
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PyPDF2 no está instalado. Ejecuta: pip install PyPDF2"
        )
    except Exception as e:
        logger.error(f"Error extrayendo texto de PDF: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Error al leer archivo PDF: {str(e)}"
        )

async def convert_document(input_path: Path, output_path: Path, output_format: str):
    """Convierte documentos con manejo robusto de errores"""
    input_ext = input_path.suffix.lower()
    
    # Lista de encodings a probar
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'windows-1252']
    
    try:
        if output_format == "txt":
            # Conversión a texto plano
            if input_ext == ".pdf":
                try:
                    text_content = extract_text_from_pdf(input_path)
                    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                        await f.write(text_content)
                except HTTPException:
                    raise
                except Exception as e:
                    logger.error(f"Error convirtiendo PDF a TXT: {str(e)}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error al convertir PDF a TXT: {str(e)}"
                    )
            elif input_ext in [".docx"]:
                try:
                    from docx import Document
                    doc = Document(input_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                        await f.write(text)
                except Exception as e:
                    logger.error(f"Error leyendo DOCX: {str(e)}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error al leer archivo DOCX: {str(e)}"
                    )
            elif input_ext in [".html", ".htm"]:
                # Convertir HTML a texto plano (básico)
                try:
                    import html
                    async with aiofiles.open(input_path, 'r', encoding='utf-8') as f:
                        html_content = await f.read()
                    # Remover tags HTML básicos
                    text = html.unescape(html_content)
                    # Remover tags HTML simples (muy básico)
                    import re
                    text = re.sub(r'<[^>]+>', '', text)
                    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                        await f.write(text)
                except Exception as e:
                    # Si falla, intentar como texto plano
                    for encoding in encodings:
                        try:
                            async with aiofiles.open(input_path, 'r', encoding=encoding) as f:
                                content = await f.read()
                            async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                                await f.write(content)
                            break
                        except:
                            continue
                    else:
                        raise HTTPException(status_code=500, detail="Error al leer el archivo")
            else:
                # Leer como texto plano con múltiples encodings
                content = None
                for encoding in encodings:
                    try:
                        async with aiofiles.open(input_path, 'r', encoding=encoding) as f:
                            content = await f.read()
                        break
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                    except Exception as e:
                        logger.error(f"Error leyendo archivo con encoding {encoding}: {str(e)}")
                        continue
                
                if content is None:
                    raise HTTPException(
                        status_code=500,
                        detail="No se pudo leer el archivo con ningún encoding compatible"
                    )
                
                async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                    await f.write(content)
        
        elif output_format == "html":
            # Conversión a HTML
            if input_ext == ".pdf":
                try:
                    text_content = extract_text_from_pdf(input_path)
                    # Escapar HTML
                    import html
                    escaped_text = html.escape(text_content)
                    # Convertir saltos de línea a <br>
                    escaped_text = escaped_text.replace('\n', '<br>\n')
                    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Documento Convertido</title>
</head>
<body>
    <pre>{escaped_text}</pre>
</body>
</html>"""
                    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                        await f.write(html_content)
                except HTTPException:
                    raise
                except Exception as e:
                    logger.error(f"Error convirtiendo PDF a HTML: {str(e)}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error al convertir PDF a HTML: {str(e)}"
                    )
            elif input_ext in [".docx"]:
                try:
                    from docx import Document
                    doc = Document(input_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    # Escapar HTML
                    import html
                    text = html.escape(text)
                    # Convertir saltos de línea a <br>
                    text = text.replace('\n', '<br>\n')
                    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Documento Convertido</title>
</head>
<body>
    <pre>{text}</pre>
</body>
</html>"""
                    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                        await f.write(html_content)
                except Exception as e:
                    logger.error(f"Error convirtiendo DOCX a HTML: {str(e)}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error al convertir DOCX a HTML: {str(e)}"
                    )
            else:
                # Leer archivo y envolver en HTML
                content = None
                for encoding in encodings:
                    try:
                        async with aiofiles.open(input_path, 'r', encoding=encoding) as f:
                            content = await f.read()
                        break
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                    except Exception as e:
                        logger.error(f"Error leyendo archivo: {str(e)}")
                        continue
                
                if content is None:
                    raise HTTPException(
                        status_code=500,
                        detail="No se pudo leer el archivo con ningún encoding compatible"
                    )
                
                # Escapar HTML
                import html
                escaped_content = html.escape(content)
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Documento Convertido</title>
</head>
<body>
    <pre>{escaped_content}</pre>
</body>
</html>"""
                async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                    await f.write(html_content)
        
        elif output_format == "md":
            # Conversión a Markdown
            if input_ext == ".pdf":
                try:
                    text_content = extract_text_from_pdf(input_path)
                    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                        await f.write(text_content)
                except HTTPException:
                    raise
                except Exception as e:
                    logger.error(f"Error convirtiendo PDF a MD: {str(e)}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error al convertir PDF a Markdown: {str(e)}"
                    )
            elif input_ext in [".docx"]:
                try:
                    from docx import Document
                    doc = Document(input_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                        await f.write(text)
                except Exception as e:
                    logger.error(f"Error convirtiendo DOCX a MD: {str(e)}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error al convertir DOCX a Markdown: {str(e)}"
                    )
            else:
                # Leer como texto y guardar como MD
                content = None
                for encoding in encodings:
                    try:
                        async with aiofiles.open(input_path, 'r', encoding=encoding) as f:
                            content = await f.read()
                        break
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                
                if content is None:
                    raise HTTPException(
                        status_code=500,
                        detail="No se pudo leer el archivo con ningún encoding compatible"
                    )
                
                async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
                    await f.write(content)
        
        elif output_format == "pdf":
            # Conversión a PDF
            try:
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import inch
                
                # Leer contenido del archivo
                text_content = ""
                
                if input_ext in [".docx"]:
                    try:
                        from docx import Document
                        doc = Document(input_path)
                        text_content = "\n".join([para.text for para in doc.paragraphs])
                    except Exception as e:
                        logger.error(f"Error leyendo DOCX para PDF: {str(e)}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Error al leer archivo DOCX: {str(e)}"
                        )
                elif input_ext == ".pdf":
                    raise HTTPException(
                        status_code=400,
                        detail="El archivo ya es un PDF. No es necesario convertirlo."
                    )
                else:
                    # Leer como texto
                    content = None
                    for encoding in encodings:
                        try:
                            async with aiofiles.open(input_path, 'r', encoding=encoding) as f:
                                content = await f.read()
                            break
                        except (UnicodeDecodeError, UnicodeError):
                            continue
                    
                    if content is None:
                        raise HTTPException(
                            status_code=500,
                            detail="No se pudo leer el archivo con ningún encoding compatible"
                        )
                    text_content = content
                
                # Crear PDF
                c = canvas.Canvas(str(output_path), pagesize=A4)
                width, height = A4
                
                # Configuración de texto
                y_position = height - 50
                line_height = 14
                margin = 50
                max_width = width - (2 * margin)
                
                # Dividir texto en líneas que quepan en la página
                lines = text_content.split('\n')
                for line in lines:
                    # Si la línea es muy larga, dividirla
                    words = line.split(' ')
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        text_width = c.stringWidth(test_line, "Helvetica", 10)
                        
                        if text_width > max_width and current_line:
                            # Dibujar línea actual
                            c.drawString(margin, y_position, current_line)
                            y_position -= line_height
                            current_line = word
                            
                            # Nueva página si es necesario
                            if y_position < margin:
                                c.showPage()
                                y_position = height - 50
                        else:
                            current_line = test_line
                    
                    # Dibujar última línea
                    if current_line:
                        c.drawString(margin, y_position, current_line)
                        y_position -= line_height
                    
                    # Nueva página si es necesario
                    if y_position < margin:
                        c.showPage()
                        y_position = height - 50
                
                c.save()
                
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="Librería reportlab no está instalada. Ejecuta: pip install reportlab"
                )
            except Exception as e:
                logger.error(f"Error creando PDF: {str(e)}\n{traceback.format_exc()}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Error al crear PDF: {str(e)}"
                )
        
        elif output_format == "docx":
            # Conversión a DOCX
            try:
                from docx import Document
                
                # Leer contenido del archivo
                text_content = ""
                
                if input_ext in [".docx"]:
                    raise HTTPException(
                        status_code=400,
                        detail="El archivo ya es un DOCX. No es necesario convertirlo."
                    )
                elif input_ext == ".pdf":
                    try:
                        text_content = extract_text_from_pdf(input_path)
                    except HTTPException:
                        raise
                    except Exception as e:
                        logger.error(f"Error leyendo PDF para DOCX: {str(e)}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Error al leer archivo PDF: {str(e)}"
                        )
                else:
                    # Leer como texto
                    content = None
                    for encoding in encodings:
                        try:
                            async with aiofiles.open(input_path, 'r', encoding=encoding) as f:
                                content = await f.read()
                            break
                        except (UnicodeDecodeError, UnicodeError):
                            continue
                    
                    if content is None:
                        raise HTTPException(
                            status_code=500,
                            detail="No se pudo leer el archivo con ningún encoding compatible"
                        )
                    text_content = content
                
                # Crear documento DOCX
                doc = Document()
                
                # Dividir en párrafos (por líneas vacías o saltos de línea)
                paragraphs = text_content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Dividir líneas largas
                        lines = para_text.split('\n')
                        for line in lines:
                            if line.strip():
                                doc.add_paragraph(line.strip())
                        # Agregar espacio entre párrafos
                        if para_text != paragraphs[-1]:
                            doc.add_paragraph()
                
                # Guardar documento
                doc.save(str(output_path))
                
            except Exception as e:
                logger.error(f"Error creando DOCX: {str(e)}\n{traceback.format_exc()}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Error al crear DOCX: {str(e)}"
                )
        
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Conversión de documentos a {output_format} no está implementada aún. Formatos disponibles: txt, html, md, pdf, docx"
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en convert_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Error al convertir documento: {str(e)}"
        )

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Descarga el archivo convertido"""
    file_path = OUTPUT_DIR / filename
    
    if not file_path.exists():
        logger.warning(f"Intento de descargar archivo no encontrado: {filename}")
        raise HTTPException(status_code=404, detail="Archivo no encontrado. El archivo puede haber expirado. Por favor, convierte el archivo nuevamente.")
    
    # Verificar que el archivo no esté vacío
    if file_path.stat().st_size == 0:
        logger.error(f"Archivo vacío detectado: {filename}")
        raise HTTPException(status_code=500, detail="El archivo está vacío o corrupto")
    
    # Obtener el nombre original del archivo si es posible
    # El filename tiene formato: {file_id}.{extension}
    try:
        # Intentar obtener el tipo MIME correcto
        extension = filename.split('.')[-1].lower()
        mime_types = {
            'mp3': 'audio/mpeg',
            'wav': 'audio/wav',
            'aac': 'audio/aac',
            'ogg': 'audio/ogg',
            'flac': 'audio/flac',
            'mp4': 'video/mp4',
            'avi': 'video/x-msvideo',
            'mov': 'video/quicktime',
            'webm': 'video/webm',
            'mkv': 'video/x-matroska',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'png': 'image/png',
            'gif': 'image/gif',
            'webp': 'image/webp',
            'bmp': 'image/bmp',
            'pdf': 'application/pdf',
            'txt': 'text/plain',
            'html': 'text/html',
            'md': 'text/markdown',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        media_type = mime_types.get(extension, 'application/octet-stream')
    except:
        media_type = 'application/octet-stream'
    
    logger.info(f"Descargando archivo: {filename} (tamaño: {file_path.stat().st_size} bytes)")
    
    return FileResponse(
        file_path,
        media_type=media_type,
        filename=filename,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        }
    )

@app.delete("/cleanup/{filename}")
async def cleanup_file(filename: str):
    """Elimina archivos temporales"""
    upload_file = UPLOAD_DIR / filename
    output_file = OUTPUT_DIR / filename
    
    if upload_file.exists():
        upload_file.unlink()
    if output_file.exists():
        output_file.unlink()
    
    return {"success": True, "message": "Archivos eliminados"}

if __name__ == "__main__":
    import uvicorn
    # Fly.io usa la variable de entorno PORT, si no existe usar 8000
    port = int(os.getenv("PORT", 8000))
    local_ip = get_local_ip()
    print("\n" + "="*60)
    print("🚀 Convertidor de Archivos - Servidor iniciado")
    print("="*60)
    print(f"📍 Acceso local:    http://localhost:{port}")
    print(f"🌐 Acceso en red:    http://{local_ip}:{port}")
    print(f"📚 Documentación:    http://localhost:{port}/docs")
    print("="*60 + "\n")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")

